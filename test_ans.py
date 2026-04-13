import httpx, json, struct, io, time
from pathlib import Path

BASE = "http://localhost:8000"
USER_ID = "test_user"
TIMEOUT = 300  # 5 минут для тяжёлых моделей

def ok(label, resp, elapsed=None):
    icon = '✅' if resp.status_code < 400 else '❌'
    t = f' ({elapsed:.1f}s)' if elapsed else ''
    print(f"{icon} [{resp.status_code}] {label}{t}")
    try: print(json.dumps(resp.json(), ensure_ascii=False, indent=2))
    except: print(resp.text)
    print()

def timed_get(url, **kw):
    t0 = time.time()
    r = httpx.get(url, timeout=TIMEOUT, **kw)
    return r, time.time() - t0

def timed_post(url, method="POST", **kw):
    t0 = time.time()
    r = getattr(httpx, method.lower())(url, timeout=TIMEOUT, **kw)
    return r, time.time() - t0

# Создаём тестовые файлы
Path('/tmp/test_doc.txt').write_text('Это тестовый документ. Python — мощный язык. FastAPI быстрый.', encoding='utf-8')

from docx import Document
doc = Document()
doc.add_heading('Тестовый документ', level=1)
doc.add_paragraph('Python — мощный язык. MTS AI Hub — платформа для ИИ.')
doc.save('/tmp/test_doc.docx')

pdf = b'%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n4 0 obj<</Length 44>>stream\nBT /F1 12 Tf 72 700 Td (Test PDF) Tj ET\nendstream\nendobj\n5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\nxref\n0 6\n0000000000 65535 f \ntrailer<</Size 6/Root 1 0 R>>\nstartxref\n0\n%%EOF'
Path('/tmp/test_doc.pdf').write_bytes(pdf)

sr = 16000
wav_data = b'\x00\x00' * sr
wav = io.BytesIO()
wav.write(b'RIFF')
wav.write(struct.pack('<I', 36 + len(wav_data)))
wav.write(b'WAVEfmt ')
wav.write(struct.pack('<IHHIIHH', 16, 1, 1, sr, sr*2, 2, 16))
wav.write(b'data')
wav.write(struct.pack('<I', len(wav_data)))
wav.write(wav_data)
Path('/tmp/test_audio.wav').write_bytes(wav.getvalue())

print('✅ Файлы: test_doc.txt, test_doc.pdf, test_doc.docx, test_audio.wav')

r, t = timed_post(f"{BASE}/v1/chat/completions", json={
    "model": "auto",
    "messages": [{"role": "user", "content": "Напиши калькулятор на Python с графическим интерфейсом"}],
    "stream": False, "user": USER_ID
})
ok("Chat / Auto Router → код", r, t)